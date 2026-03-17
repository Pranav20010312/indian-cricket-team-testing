"""
Generate Testing_Approach_Document.docx using python-docx.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ---- Page setup ----
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ---- Styles ----
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

h1 = doc.styles["Heading 1"]
h1.font.name = "Arial"
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(10)

h2 = doc.styles["Heading 2"]
h2.font.name = "Arial"
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
h2.paragraph_format.space_before = Pt(14)
h2.paragraph_format.space_after = Pt(8)


def add_para(text, bold=False, italic=False, font_name=None, font_size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name or "Arial"
    run.font.size = Pt(font_size or 11)
    run.bold = bold
    run.italic = italic
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return p


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_test_table(rows_data):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(2.8)
        row.cells[2].width = Inches(1.7)

    # Header row
    headers = ["Test Name", "What It Tests", "Expected Result"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1F3864")

    # Data rows
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = "Arial"
        if idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F2F7FB")

    return table


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Testing Approach Document")
run.font.size = Pt(28)
run.bold = True
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Indian Cricket Team Web Application")
run.font.size = Pt(16)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by: Pranav Ramamurthy")
run.font.size = Pt(12)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared for: Qrvey (Amit Bhatnagar, Olga Lisnichenko)")
run.font.size = Pt(12)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Date: March 2026")
run.font.size = Pt(12)
run.font.name = "Arial"
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_page_break()

# ============================================================
# SECTION 1: PROJECT OVERVIEW
# ============================================================
doc.add_heading("1. Project Overview", level=1)

add_para(
    "This project demonstrates a complete approach to web application development, "
    "testing, and intelligent regression detection. The deliverable is a Flask-based "
    "web application themed around the Indian Cricket Team, accompanied by a traditional "
    "automated test suite and an AI-powered regression testing system."
)

add_para(
    "The application was built specifically for the Qrvey assessment to address a real "
    "challenge: automating regression testing without relying on manually written test cases. "
    "Qrvey has already automated much of their testing with AI but struggles with regression "
    "testing, where test cases become outdated as features evolve. This project presents a "
    "working proof-of-concept that solves that problem."
)

add_para("Tech stack and rationale:", bold=True, space_after=4)
add_bullet("Python Flask: Lightweight, minimal boilerplate, fast to prototype. Perfect for demonstrating concepts without framework overhead.")
add_bullet("Jinja2 + HTML/CSS: Server-rendered templates keep the architecture simple. No JavaScript framework needed for this scope.")
add_bullet("Selenium + pytest: Industry-standard browser automation. Pytest provides clean fixtures, parameterization, and HTML reporting.")
add_bullet("Anthropic Claude API: Powers the AI analysis layer that explains regression findings in plain English with root cause analysis.")
add_bullet("BeautifulSoup + requests: Used by the regression system to crawl and parse the running application without a browser.")

# ============================================================
# SECTION 2: APPLICATION ARCHITECTURE
# ============================================================
doc.add_heading("2. Application Architecture", level=1)

add_para(
    "The application follows a standard Flask server-rendered architecture. A single app.py "
    "file serves as the backend, handling routing, authentication, and data. Templates are "
    "rendered server-side with Jinja2, and static CSS provides the Indian cricket team theme "
    "(blue #0033A0 and orange #FF6B00)."
)

add_para("Login flow:", bold=True, space_after=4)
add_para(
    "The user submits credentials via a POST request to /login. The server validates against "
    "hardcoded values (admin / admin123). On success, the username is stored in Flask\u2019s "
    "session object and the user is redirected to /dashboard. On failure, an error message is "
    "rendered on the same page with specific feedback (wrong credentials, empty fields)."
)

add_para("Route protection:", bold=True, space_after=4)
add_para(
    "A login_required decorator wraps every protected route. It checks for a username key in "
    "the Flask session. If absent, it redirects to /login. This ensures /dashboard, /players, "
    "/matches, /api/players, and /api/matches are all inaccessible without authentication. "
    "Logout clears the session entirely via session.clear() and redirects to the login page."
)

add_para("Data layer:", bold=True, space_after=4)
add_para(
    "Player and match data are stored as Python lists of dictionaries directly in app.py. "
    "This keeps the demo self-contained with no database setup required. The same data serves "
    "both the HTML templates (via Jinja2 context) and the JSON API endpoints (/api/players, "
    "/api/matches)."
)

doc.add_page_break()

# ============================================================
# SECTION 3: TRADITIONAL TESTING APPROACH
# ============================================================
doc.add_heading("3. Traditional Testing Approach", level=1)

add_para(
    "The test suite follows a three-layer strategy that covers the application from the inside out:"
)

add_para("Layer 1 \u2013 API level (test_api.py):", bold=True, space_after=4)
add_para(
    "Tests the JSON endpoints directly using Python\u2019s requests library. This is the fastest "
    "feedback loop. It validates status codes, response structure, field presence, and "
    "authentication enforcement. No browser needed, so these tests run in milliseconds."
)

add_para("Layer 2 \u2013 UI level (test_login.py, test_dashboard.py, test_players.py, test_matches.py):", bold=True, space_after=4)
add_para(
    "Tests the application through a real browser using Selenium WebDriver in headless Chrome. "
    "These verify what the user actually sees: form submissions, page navigation, table rendering, "
    "search filtering, sorting, and button interactions. Each test gets a fresh browser instance to "
    "prevent state leakage."
)

add_para("Layer 3 \u2013 Security level (test_security.py):", bold=True, space_after=4)
add_para(
    "Tests that unauthenticated users cannot access protected routes. Verifies that direct URL "
    "access to /dashboard, /players, and /matches redirects to /login, and that logging out "
    "properly invalidates the session."
)

add_para("Why Selenium + pytest:", bold=True, space_after=4)
add_para(
    "Selenium is the de facto standard for browser automation testing. It drives a real browser, "
    "catching rendering issues that HTTP-only tests miss. Pytest was chosen over unittest for its "
    "fixture system (conftest.py provides shared setup), automatic test discovery, and the "
    "pytest-html plugin for generating visual reports. The entire suite runs with a single command: "
    "pytest tests/ -v --html=report.html."
)

add_para("Test organization:", bold=True, space_after=4)
add_para(
    "Each test file is independent and focuses on one area of the application. A shared conftest.py "
    "starts the Flask server in a background thread, provides browser fixtures, and handles cleanup. "
    "The logged_in_browser fixture eliminates repetitive login code across tests. Total: 37 automated "
    "tests across 6 files."
)

doc.add_page_break()

# ============================================================
# SECTION 4: AI-POWERED REGRESSION TESTING APPROACH
# ============================================================
doc.add_heading("4. AI-Powered Regression Testing Approach", level=1)

doc.add_heading("The Problem", level=2)
add_para(
    "Manual regression testing has a fundamental scaling problem. Every new feature requires "
    "updating existing test cases to account for changed behavior. Over time, test suites "
    "accumulate hundreds of assertions that become brittle and outdated. When a test fails, "
    "developers spend more time figuring out whether the test is wrong or the code is wrong. "
    "Qrvey has experienced this firsthand: their testing is largely automated with AI, but "
    "regression testing remains a pain point because test cases need constant maintenance."
)

doc.add_heading("The Solution", level=2)
add_para(
    "This project demonstrates a different approach: baseline comparison combined with AI analysis. "
    "Instead of writing test cases that assert specific values, the system learns what the "
    "application looks like when it is working correctly, then detects any deviation after a code "
    "change. No test cases to write. No test cases to maintain."
)

add_para("The system has four components:", space_after=4)

add_para("Baseline Capture (capture_baseline.py):", bold=True, space_after=4)
add_para(
    "Crawls the running application and records a complete snapshot. For each route, it captures "
    "the HTTP status code, HTML element IDs and class names, heading text, table row counts, card "
    "counts, content length, and response time. For API endpoints, it records the JSON schema "
    "(top-level keys, array item fields, counts). Everything is saved as baseline.json."
)

add_para("Regression Detector (detect_regression.py):", bold=True, space_after=4)
add_para(
    "After a code change, this script crawls the same routes and compares each response against "
    "the baseline. It checks for status code changes (critical), missing HTML elements or JSON "
    "fields (high), data count changes (medium), response time increases over 2x (high), and new "
    "unexpected elements (low/warning). Each finding is classified as REGRESSION (unexpected "
    "change), WARNING (needs review), or EXPECTED (developer-marked known change)."
)

add_para("AI Analyzer (ai_analyzer.py):", bold=True, space_after=4)
add_para(
    "Takes the regression report and sends it to Claude via the Anthropic API. The AI reads every "
    "finding and provides a plain English explanation, severity classification with reasoning, "
    "probable root cause, recommended fix, and whether the finding is likely intentional or an "
    "actual regression. If no API key is configured, the script falls back to rule-based analysis "
    "so the demo always works."
)

add_para("Simulation Script (simulate_change.py):", bold=True, space_after=4)
add_para(
    "Intentionally introduces regressions to demonstrate the system. It removes a player from the "
    "data (changes row count and API count), drops a field from one player\u2019s JSON (schema "
    "change), and adds a 3-second delay to an endpoint (performance regression). After the demo, "
    "it reverts all changes from a backup."
)

doc.add_heading("How It Flows", level=2)
flow_steps = [
    "Start the application and capture the baseline (the known-good state)",
    "Make a code change (or run the simulation script)",
    "Run the regression detector to compare current state against baseline",
    "Run the AI analyzer to get intelligent explanations and fix recommendations",
    "Review the report and fix actual regressions; dismiss expected changes",
]
for i, step in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {step}")
    run.font.name = "Arial"
    run.font.size = Pt(11)

doc.add_heading("How This Scales", level=2)
add_para(
    "The key advantage is zero maintenance. When a developer adds a new feature, the baseline is "
    "recaptured. The next time a change is made, the system automatically checks whether the new "
    "feature still works alongside everything else. There are no test cases to update. The system "
    "simply compares what is to what was and flags anything that changed unexpectedly. As the "
    "application grows, the baseline grows with it automatically."
)

doc.add_page_break()

# ============================================================
# SECTION 5: TEST CASES SUMMARY
# ============================================================
doc.add_heading("5. Test Cases Summary", level=1)

add_para("Login Tests (8 tests)", bold=True, space_after=4)
add_test_table([
    ["Successful login", "Correct credentials redirect to dashboard", "Dashboard loads"],
    ["Wrong password", "Invalid password shows error", "Error message displayed"],
    ["Wrong username", "Invalid username shows error", "Error message displayed"],
    ["Empty username", "Missing username validation", "Username error shown"],
    ["Empty password", "Missing password validation", "Password error shown"],
    ["Both fields empty", "Empty form submission", "Validation error shown"],
    ["Error msg visible", "Error element renders on failure", "Error div is displayed"],
    ["Redirect after login", "Post-login landing page", "URL contains /dashboard"],
])
doc.add_paragraph()

add_para("Dashboard Tests (8 tests)", bold=True, space_after=4)
add_test_table([
    ["Dashboard loads", "Page loads after authentication", "Dashboard content visible"],
    ["Welcome message", "Username shown in greeting", "Welcome, admin displayed"],
    ["View Players btn", "Button exists and is clickable", "Element present and enabled"],
    ["View Matches btn", "Button exists and is clickable", "Element present and enabled"],
    ["View Stats btn", "Button exists and is clickable", "Element present and enabled"],
    ["Logout works", "Logout clears session", "Redirects to /login"],
    ["Players nav", "Clicking navigates to players", "URL contains /players"],
    ["Matches nav", "Clicking navigates to matches", "URL contains /matches"],
])
doc.add_paragraph()

add_para("Players Page Tests (7 tests)", bold=True, space_after=4)
add_test_table([
    ["Table loads", "Players table renders with rows", "At least 8 rows present"],
    ["Correct columns", "Table has all 5 column headers", "All 5 headers found"],
    ["Search by name", "Typing a name filters rows", "Only matching rows visible"],
    ["Search by role", "Typing a role filters rows", "Matching role rows shown"],
    ["No results msg", "Searching nonsense text", "No results message displayed"],
    ["Back button", "Back to Dashboard navigation", "URL contains /dashboard"],
    ["Sort columns", "Clicking header sorts data", "Rows reordered correctly"],
])
doc.add_paragraph()

add_para("Match Scores Tests (4 tests)", bold=True, space_after=4)
add_test_table([
    ["Results load", "Match cards render with data", "At least 5 cards present"],
    ["Filter wins", "Wins button filters results", "Only won matches visible"],
    ["Filter losses", "Losses button filters results", "Only lost matches visible"],
    ["Back button", "Back to Dashboard navigation", "URL contains /dashboard"],
])
doc.add_paragraph()

add_para("Security Tests (4 tests)", bold=True, space_after=4)
add_test_table([
    ["/dashboard redirect", "Unauthenticated access blocked", "Redirects to /login"],
    ["/players redirect", "Unauthenticated access blocked", "Redirects to /login"],
    ["/matches redirect", "Unauthenticated access blocked", "Redirects to /login"],
    ["Post-logout redirect", "Session cleared on logout", "Dashboard redirects to /login"],
])
doc.add_paragraph()

add_para("API Tests (6 tests)", bold=True, space_after=4)
add_test_table([
    ["GET /api/players 200", "Authenticated API returns data", "Status 200, valid JSON"],
    ["GET /api/matches 200", "Authenticated API returns data", "Status 200, valid JSON"],
    ["Players structure", "JSON has correct fields", "name, role, matches, runs, wickets"],
    ["Matches structure", "JSON has correct fields", "opponent, date, venue, result, score"],
    ["Players no auth", "Unauthenticated API access", "Returns 302 redirect"],
    ["Matches no auth", "Unauthenticated API access", "Returns 302 redirect"],
])

doc.add_page_break()

# ============================================================
# SECTION 6: HOW TO RUN EVERYTHING
# ============================================================
doc.add_heading("6. How to Run Everything", level=1)

add_para("Prerequisites: Python 3.8+, Google Chrome, ChromeDriver.")
doc.add_paragraph()

add_para("Step 1: Install dependencies", bold=True, space_after=4)
add_para("pip install -r requirements.txt", font_name="Consolas", font_size=10)
doc.add_paragraph()

add_para("Step 2: Start the Flask application", bold=True, space_after=4)
add_para("python app.py", font_name="Consolas", font_size=10)
add_para("The app runs at http://127.0.0.1:5000. Login with admin / admin123.")
doc.add_paragraph()

add_para("Step 3: Run the automated test suite (in a second terminal)", bold=True, space_after=4)
add_para("pytest tests/ -v --html=report.html --self-contained-html", font_name="Consolas", font_size=10)
add_para("This runs all 37 tests and generates an HTML report.")
doc.add_paragraph()

add_para("Step 4: Run the AI regression testing demo", bold=True, space_after=4)
add_para("python ai_regression/capture_baseline.py", font_name="Consolas", font_size=10)
add_para("python ai_regression/simulate_change.py --apply", font_name="Consolas", font_size=10)
add_para("python ai_regression/detect_regression.py", font_name="Consolas", font_size=10)
add_para("python ai_regression/ai_analyzer.py", font_name="Consolas", font_size=10)
add_para("python ai_regression/simulate_change.py --revert", font_name="Consolas", font_size=10)
doc.add_paragraph()
add_para(
    "Set the ANTHROPIC_API_KEY environment variable before running the AI analyzer for "
    "Claude-powered analysis. Without it, the script uses a rule-based fallback."
)

# ============================================================
# SECTION 7: HOW THIS CONNECTS TO QRVEY
# ============================================================
doc.add_heading("7. How This Connects to Qrvey", level=1)

add_para(
    "Qrvey has automated most of their testing with AI but struggles specifically with regression "
    "testing. The core issue is that traditional regression test cases require constant maintenance "
    "as the product evolves. Every feature change risks breaking existing tests, and every broken "
    "test requires investigation to determine whether the test is outdated or the code is actually "
    "broken."
)

add_para(
    "The baseline comparison approach demonstrated in this project addresses that directly. There "
    "are no test cases to maintain. The system captures what the application looks like when it "
    "works, and after any change, it automatically detects deviations. The developer does not need "
    "to predict what might break or write assertions for every possible scenario."
)

add_para(
    "The AI analysis layer adds a dimension that raw test automation cannot provide. Instead of a "
    "binary pass/fail result, the AI explains why something changed, classifies its severity, "
    "identifies the probable root cause, and recommends a fix. This reduces the time developers "
    "spend investigating failures."
)

add_para(
    "This approach extends naturally to Qrvey\u2019s specific challenges. Multi-tenant security "
    "testing: capture a baseline of what Tenant A sees, make a change, and verify Tenant A\u2019s "
    "data has not leaked to Tenant B. Localization: capture baselines across languages and detect "
    "when a translation breaks after a code change. API contract stability: capture JSON schemas "
    "and flag any breaking change before it reaches consumers."
)

add_para(
    "Future direction: the system could be made fully agentic using a framework like LangGraph, "
    "where the AI not only detects and explains regressions but actively investigates them by "
    "trying different inputs, narrowing down the root cause, and even suggesting code patches. "
    "The proof-of-concept in this project is the foundation for that vision."
)

# ---- Save ----
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "Testing_Approach_Document.docx"
)
doc.save(output_path)
print(f"Document created: {output_path}")
